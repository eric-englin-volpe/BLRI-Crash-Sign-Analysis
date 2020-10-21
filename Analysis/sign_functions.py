

def clean_sign_data(df):
    df_clean = df.loc[df['Sign Type']!="Not a sign on the road"]
    df_clean = df_clean.loc[df_clean['Sign Type']!="Flag for Removal"]
    bad_signs_count = 6692-len(df_clean)
    print(bad_signs_count, " pictures signs flagged to remove")

    df_clean = df_clean.drop_duplicates(subset = ['Physical Location', 'Parkway Location',
                                            'Mile Post (tenth of a mile)', 'Specific Sign'])
    dups_count = (6692- bad_signs_count) - len(df_clean)
    print(dups_count, " duplicates")
    df_clean = df_clean.loc[df_clean['distance (miles)']<1]
    print((6692-dups_count-bad_signs_count) - len(df_clean), " more than a mile from the road")
    print(len(df_clean), " in clean dataset")
    #df_clean = df_clean.loc[df_clean['Sign Type']==type]
    #if subtype != "all":
#        df_clean = df_clean.loc[df_clean['Sign Subtype']==subtype]
#    if specific != "all":
#        df_clean = df_clean.loc[df_clean['Specific Sign']==specific]
#    print(len(df_clean), " in analysis dataset")
    return df_clean

def make_location_plot(df, title, filter_name, folder_name):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import numpy as np
    import seaborn as sns
    import warnings
    warnings.filterwarnings("ignore")
    graph_title = filter_name + " " + title + " Locations on BLRI Parkway"
    axis_title = title + " Count"
    plt.style.use('seaborn')
    fig, ax = plt.subplots(1,1,figsize=(20,10))
    tick_spacing = 50

    plt.bar(df['Mile Post (ten miles)'].value_counts().index+5, df['Mile Post (ten miles)'].value_counts().values,
            width = 10)

    bar_max = df['Mile Post (ten miles)'].value_counts().max()*1.05
    plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)

    if title == "Sign":
        ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 0.25)

        plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
        bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
        ax.text(279, bar_max/2,
                   '  No Data\n  Collected\n  MP 276-315', fontsize=13,
               bbox=bbox_props)
    else:
        ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 2)

    ax.text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=15)

    ax.text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=15)
    ax.text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=15)
    ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
    ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
    #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
    #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax.set_ylim([0, bar_max*1.3])
    ax.set_xlim([0, 470])
#    ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
    ax.set_title(" ",fontsize= 12, pad = 20) # title of plot

    ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax.set_ylabel(axis_title, fontsize = 30)#ylabel
    ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax.tick_params(axis='x', which='major', labelsize=16)
    ax.tick_params(axis='y', which='major', labelsize=16)
    ax.tick_params(axis='both', which='minor', labelsize=8)
    ax.grid('off')
    plt.legend(fontsize = 20)
    file_save_name = folder_name + "/" + title + "_locations.png"
    #print(file_save_name)
    plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
    plt.close()

    ############################################################################
    #sign type plot
    if title == "Sign":
        for sign_type in ['Regulatory','Warning','Recreational','Guide','General Information','Others', 'Others 2']:
            tot_signs_list= []
            ten_mile_markers_for_graph = []
            for x in range(0, 470, 10):
                ten_mile_markers_for_graph.append(x+5)
                temp_df = df.loc[df['Sign Type']==sign_type]
                if sign_type == "Others":
                    temp_df = df.loc[df['Sign Type'].isin(['Regulatory','Warning'])==False]
                if sign_type == "Others 2":
                    temp_df = df.loc[df['Sign Type'].isin(['Misc. Other','General Service Signs'])]
                temp_df = temp_df.loc[temp_df['Mile Post (ten miles)']==x]
                tot_signs = temp_df['Sign Type'].count()
                tot_signs_list.append(tot_signs)
            if sign_type == "Others":
                others_list = tot_signs_list
            if sign_type == "Regulatory":
                regulatory_list = tot_signs_list
            if sign_type == "Warning":
                warning_list = tot_signs_list
            if sign_type == "Recreational":
                recreational_list = tot_signs_list
            if sign_type == "Guide":
                guide_list = tot_signs_list
            if sign_type == "General Information":
                general_information_list = tot_signs_list
            if sign_type == "Others 2":
                others2_list = tot_signs_list

        #top 2 sign type graph
        graph_title = filter_name + " " + title + " by Sign Type Locations on BLRI Parkway"
        axis_title = title + " Count"
        plt.style.use('seaborn')
        sns.set_style("ticks")
        fig, ax = plt.subplots(1,1,figsize=(20,15))
        tick_spacing = 50

        dataset1 = np.array(regulatory_list)
        dataset2 = np.array(warning_list)
        dataset3 = np.array(others_list)

        p1 = plt.bar(ten_mile_markers_for_graph, dataset1, width=10)
        p2 = plt.bar(ten_mile_markers_for_graph, dataset2,  width=10, bottom=dataset1, color="salmon",hatch = "/")
        p3 = plt.bar(ten_mile_markers_for_graph, dataset3,  width=10, bottom=dataset1+dataset2,color = 'gray')
        plt.legend(['Regulatory','Warning','All Other Sign Types'], fontsize = 20)

        bar_max = max(dataset1+dataset2+dataset3)*1.05
        plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
        plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
        plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)

        ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 0.25)

        plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
        bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
        ax.text(279, bar_max/2,
                   '  No Data\n  Collected\n  MP 276-315', fontsize=13,
               bbox=bbox_props)


        ax.text(95, bar_max,
                   'Roanoke, VA\n MP 100-130', fontsize=15)

        ax.text(275,bar_max,
                   'Boone, NC\n MP 280-300', fontsize=15)
        ax.text(365, bar_max,
                   'Asheville,NC\n MP 370-400', fontsize=15)
        ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 2)
        ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                    linestyles = "dotted", linewidth = 5, color = 'black')
        ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
        ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
        #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
        #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
        #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

        #make graph pretty
        ax.set_ylim([0, bar_max*1.5])
        ax.set_xlim([0, 470])
#        ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
        ax.set_title("",fontsize= 15, pad = 20) # title of plot

        ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
        ax.set_ylabel(axis_title, fontsize = 30)#ylabel
        ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
        ax.tick_params(axis='x', which='major', labelsize=16)
        ax.tick_params(axis='y', which='major', labelsize=16)
        ax.tick_params(axis='both', which='minor', labelsize=8)
        ax.grid('off')
        file_save_name = folder_name + "/" + "Sign Type_locations.png"
        #print(file_save_name)
        plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
        plt.close()

        #all sign type graph
        graph_title = filter_name + " " + title + " by Sign Type Locations on BLRI Parkway"
        axis_title = title + " Count"
        plt.style.use('seaborn')
        sns.set_style("ticks")
        fig, ax = plt.subplots(1,1,figsize=(20,15))
        tick_spacing = 50

        dataset1 = np.array(regulatory_list)
        dataset2 = np.array(warning_list)
        dataset3 = np.array(recreational_list)
        dataset4 = np.array(guide_list)
        dataset5 = np.array(general_information_list)
        dataset6 = np.array(others2_list)

        p1 = plt.bar(ten_mile_markers_for_graph, dataset1, width=10)
        p2 = plt.bar(ten_mile_markers_for_graph, dataset2,  width=10, bottom=dataset1, color="salmon",hatch = "/")
        p3 = plt.bar(ten_mile_markers_for_graph, dataset3,  width=10, bottom=dataset1+dataset2, color = "saddlebrown")
        p4 = plt.bar(ten_mile_markers_for_graph, dataset4,  width=10, bottom=dataset1+dataset2+dataset3, color = "forestgreen",hatch = "/")
        p5 = plt.bar(ten_mile_markers_for_graph, dataset5,  width=10, bottom=dataset1+dataset2+dataset3+dataset4, color = 'thistle')
        p6 = plt.bar(ten_mile_markers_for_graph, dataset6,  width=10, bottom=dataset1+dataset2+dataset3+dataset4+dataset5, color = 'gray',hatch = "/")
        plt.legend(['Regulatory','Warning','Recreational','Guide', 'General Information', 'Other'], ncol = 2, fontsize = 20)

        bar_max = max(dataset1+dataset2+dataset3+dataset4+dataset5+dataset6)*1.05
        plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
        plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
        plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
        ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 0.25)

        plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
        bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
        ax.text(279, bar_max/2,
                   '  No Data\n  Collected\n  MP 276-315', fontsize=13,
               bbox=bbox_props)


        ax.text(95, bar_max,
                   'Roanoke, VA\n MP 100-130', fontsize=15)

        ax.text(275,bar_max,
                   'Boone, NC\n MP 280-300', fontsize=15)
        ax.text(365, bar_max,
                   'Asheville,NC\n MP 370-400', fontsize=15)
        ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 2)
        ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                    linestyles = "dotted", linewidth = 5, color = 'black')
        ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
        ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
        #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
        #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
        #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

        #make graph pretty
        ax.set_ylim([0, bar_max*1.5])
        ax.set_xlim([0, 470])
#        ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
        ax.set_title("",fontsize= 15, pad = 20) # title of plot

        ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
        ax.set_ylabel(axis_title, fontsize = 30)#ylabel
        ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
        ax.tick_params(axis='x', which='major', labelsize=16)
        ax.tick_params(axis='y', which='major', labelsize=16)
        ax.tick_params(axis='both', which='minor', labelsize=8)
        ax.grid('off')
        file_save_name = folder_name + "/" + "Sign Type All_locations.png"
        #print(file_save_name)
        plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
        plt.close()

        ##############################################################################

        #all sign type graph - sequential colors
        graph_title = filter_name + " " + title + " by Sign Type Locations on BLRI Parkway"
        axis_title = title + " Count"
        plt.style.use('seaborn')
        sns.set_style("ticks")
        sns.set_palette("rocket", 6)
        fig, ax = plt.subplots(1,1,figsize=(20,15))
        tick_spacing = 50

        dataset1 = np.array(regulatory_list)
        dataset2 = np.array(warning_list)
        dataset3 = np.array(recreational_list)
        dataset4 = np.array(guide_list)
        dataset5 = np.array(general_information_list)
        dataset6 = np.array(others2_list)

        p1 = plt.bar(ten_mile_markers_for_graph, dataset1, width=10)
        p2 = plt.bar(ten_mile_markers_for_graph, dataset2,  width=10, bottom=dataset1,hatch = "/")
        p3 = plt.bar(ten_mile_markers_for_graph, dataset3,  width=10, bottom=dataset1+dataset2)
        p4 = plt.bar(ten_mile_markers_for_graph, dataset4,  width=10, bottom=dataset1+dataset2+dataset3,hatch = "/")
        p5 = plt.bar(ten_mile_markers_for_graph, dataset5,  width=10, bottom=dataset1+dataset2+dataset3+dataset4)
        p6 = plt.bar(ten_mile_markers_for_graph, dataset6,  width=10, bottom=dataset1+dataset2+dataset3+dataset4+dataset5,hatch = "/")
        plt.legend(['Regulatory','Warning','Recreational','Guide', 'General Information', 'Other'], fontsize = 20, ncol=2)

        bar_max = max(dataset1+dataset2+dataset3+dataset4+dataset5+dataset6)*1.05
        plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
        plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
        plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)

        ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 0.25)

        plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
        bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
        ax.text(279, bar_max/2,
                   '  No Data\n  Collected\n  MP 276-315', fontsize=13,
               bbox=bbox_props)


        ax.text(95, bar_max,
                   'Roanoke, VA\n MP 100-130', fontsize=15)

        ax.text(275,bar_max,
                   'Boone, NC\n MP 280-300', fontsize=15)
        ax.text(365, bar_max,
                   'Asheville,NC\n MP 370-400', fontsize=15)
        ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                    linestyles = "dotted", linewidth = 2)
        ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                    linestyles = "dotted", linewidth = 5, color = 'black')
        ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
        ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
        #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
        #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
        #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

        #make graph pretty
        ax.set_ylim([0, bar_max*1.5])
        ax.set_xlim([0, 470])
#        ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
        ax.set_title("",fontsize= 15, pad = 20) # title of plot

        ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
        ax.set_ylabel(axis_title, fontsize = 30)#ylabel
        ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
        ax.tick_params(axis='x', which='major', labelsize=16)
        ax.tick_params(axis='y', which='major', labelsize=16)
        ax.tick_params(axis='both', which='minor', labelsize=8)
        ax.grid('off')
        file_save_name = folder_name + "/" + "Sign Type All Sequential_locations.png"
        #print(file_save_name)
        plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
        plt.close()

        #return plt


def make_location_subtype_plots(df, title, filter_name, folder_name):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import numpy as np
    import seaborn as sns
    import warnings
    warnings.filterwarnings("ignore")

    subtype_list =  ['advance traffic control and flow',
                      'vertical alignment and other road features','weather',
                      'horizontal alignment','Others']
    for this_subtype in subtype_list:

        tot_signs_list= []
        ten_mile_markers_for_graph = []
        for x in range(0, 470, 10):
            temp_df = df.loc[df['Sign Type']=="Warning"]
            ten_mile_markers_for_graph.append(x+5)
            if this_subtype == "Others":
                temp_df = temp_df.loc[temp_df['Sign Subtype'].str.lower().isin(subtype_list)==False]
            else:
                temp_df = temp_df.loc[temp_df['Sign Subtype'].str.lower()==this_subtype]
            temp_df = temp_df.loc[temp_df['Mile Post (ten miles)']==x]
            tot_signs = temp_df['Sign Type'].count()
            tot_signs_list.append(tot_signs)

        if this_subtype == 'advance traffic control and flow':
            advance_traffic_list = tot_signs_list
        if this_subtype == 'horizontal alignment':
            horizontal_alignment_list  = tot_signs_list
        if this_subtype == 'vertical alignment and other road features':
            vert_alignment_list  = tot_signs_list
        if this_subtype == 'weather':
            weather_list  = tot_signs_list
        if this_subtype == 'Others':
            others_list  = tot_signs_list

    #print(len(advance_traffic_list), len(weather_list), len(others_list))
    graph_title = "Sign Subtypes of Warning Signs Locations on BLRI Parkway"
    axis_title = title + " Count"
    plt.style.use('seaborn')
    sns.set_style("ticks")
    sns.set_palette("pink", 5)
    fig, ax = plt.subplots(1,1,figsize=(20,15))
    tick_spacing = 50


    dataset1 = np.array(advance_traffic_list)
    dataset2 = np.array(horizontal_alignment_list)
    dataset3 = np.array(vert_alignment_list)
    dataset4 = np.array(weather_list)
    dataset5 = np.array(others_list)

    p1 = plt.bar(ten_mile_markers_for_graph, dataset1, width=10)
    p2 = plt.bar(ten_mile_markers_for_graph, dataset2,  width=10, bottom=dataset1,hatch = "/")
    p3 = plt.bar(ten_mile_markers_for_graph, dataset3,  width=10, bottom=dataset1+dataset2)
    p4 = plt.bar(ten_mile_markers_for_graph, dataset4,  width=10, bottom=dataset1+dataset2+dataset3,hatch = "/")
    p5 = plt.bar(ten_mile_markers_for_graph, dataset5,  width=10, bottom=dataset1+dataset2+dataset3+dataset4)
    plt.legend(['Advance Traffic Control','Horizontal Alignment','Vertical Alignment','Weather', 'Other'], fontsize = 20, ncol=2)

    bar_max = max(dataset1+dataset2+dataset3+dataset4+dataset5)*1.05
    plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)

    ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax.text(279, bar_max/2,
               '  No Data\n  Collected\n  MP 276-315', fontsize=13,
           bbox=bbox_props)


    ax.text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=15)

    ax.text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=15)
    ax.text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=15)
    ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
    ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
    #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
    #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax.set_ylim([0, bar_max*1.5])
    ax.set_xlim([0, 470])
#    ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
    ax.set_title(" ",fontsize= 12, pad = 20) # title of plot

    ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax.set_ylabel(axis_title, fontsize = 30)#ylabel
    ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax.tick_params(axis='x', which='major', labelsize=16)
    ax.tick_params(axis='y', which='major', labelsize=16)
    ax.tick_params(axis='both', which='minor', labelsize=8)
    ax.grid('off')
    file_save_name = folder_name + "/" + "Sign Warning Subtypes_locations.png"
    #print(file_save_name)
    plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
    plt.close()


    ##############################################################################
    #Regulatory Type Signs

    subtype_list =  ['stop',
                      'movement control and parking','exclusion',
                      'speed limit','yield','headlights']
    for this_subtype in subtype_list:

        tot_signs_list= []
        ten_mile_markers_for_graph = []
        for x in range(0, 470, 10):
            temp_df = df.loc[df['Sign Type']=="Regulatory"]
            ten_mile_markers_for_graph.append(x+5)
            temp_df = temp_df.loc[temp_df['Sign Subtype'].str.lower()==this_subtype]
            temp_df = temp_df.loc[temp_df['Mile Post (ten miles)']==x]
            tot_signs = temp_df['Sign Type'].count()
            tot_signs_list.append(tot_signs)

        if this_subtype == 'stop':
            stop_list = tot_signs_list
        if this_subtype == 'movement control and parking':
            movement_control_list  = tot_signs_list
        if this_subtype == 'exclusion':
            exclusion_list  = tot_signs_list
        if this_subtype == 'speed limit':
            speed_limit_alignment_list  = tot_signs_list
        if this_subtype == 'yield':
            yield_list  = tot_signs_list
        if this_subtype == 'headlights':
            headlights_list  = tot_signs_list

    graph_title = "Sign Subtypes of Regulatory Signs Locations on BLRI Parkway"
    axis_title = title + " Count"
    plt.style.use('seaborn')
    sns.set_style("ticks")
    sns.set_palette("mako", 6)
    fig, ax = plt.subplots(1,1,figsize=(20,15))
    tick_spacing = 50

    dataset1 = np.array(stop_list)
    dataset2 = np.array(movement_control_list)
    dataset3 = np.array(exclusion_list)
    dataset4 = np.array(speed_limit_alignment_list)
    dataset5 = np.array(yield_list)
    dataset6 = np.array(headlights_list)

    p1 = plt.bar(ten_mile_markers_for_graph, dataset1, width=10)
    p2 = plt.bar(ten_mile_markers_for_graph, dataset2,  width=10, bottom=dataset1,hatch = "/")
    p3 = plt.bar(ten_mile_markers_for_graph, dataset3,  width=10, bottom=dataset1+dataset2)
    p4 = plt.bar(ten_mile_markers_for_graph, dataset4,  width=10, bottom=dataset1+dataset2+dataset3,hatch = "/")
    p5 = plt.bar(ten_mile_markers_for_graph, dataset5,  width=10, bottom=dataset1+dataset2+dataset3+dataset4)
    p6 = plt.bar(ten_mile_markers_for_graph, dataset6,  width=10, bottom=dataset1+dataset2+dataset3+dataset4+dataset5,hatch = "/")
    plt.legend(['Stop','Movement Control','Exclusion','Speed Limit', 'Yield','Headlights'], fontsize = 20, ncol=2)

    bar_max = max(dataset1+dataset2+dataset3+dataset4+dataset5+dataset6)*1.05
    plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax.text(279, bar_max/2,
               '  No Data\n  Collected\n  MP 276-315', fontsize=13,
           bbox=bbox_props)


    ax.text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=15)

    ax.text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=15)
    ax.text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=15)
    ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
    ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)
    #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
    #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax.set_ylim([0,  bar_max*1.5])
    ax.set_xlim([0, 470])
#    ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
    ax.set_title(" ",fontsize= 12, pad = 20) # title of plot

    ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax.set_ylabel(axis_title, fontsize = 30)#ylabel
    ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax.tick_params(axis='x', which='major', labelsize=16)
    ax.tick_params(axis='y', which='major', labelsize=16)
    ax.tick_params(axis='both', which='minor', labelsize=8)
    ax.grid('off')
    file_save_name = folder_name + "/" + "Sign Regulatory Subtypes_locations.png"
    #print(file_save_name)
    plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
    plt.close()


def make_location_specific_sign_plots(df, title, filter_name, folder_name):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import numpy as np
    import seaborn as sns
    import warnings
    warnings.filterwarnings("ignore")

    mylist = df['Specific Sign'].value_counts().index.sort_values()
    if len(mylist) > 10:
        mylist = mylist[:10]

    graph_title = "Specific Signs Locations on BLRI Parkway"
    axis_title = " Count"
    plt.style.use('seaborn')
    sns.set_style("dark", {"axes.facecolor": ".9"})
    try:
        sns.set_palette("Blues_r", len(mylist)-2)
    except:
        sns.set_palette("Blues_r", len(mylist))
    fig, ax = plt.subplots(1,1,figsize=(20,15))
    tick_spacing = 50
    last_dataset= np.array(0*47)
    label_list = []

    hatch = ""
    count = 1
    for this_sign in mylist:
        this_list = []
        ten_mile_markers_for_graph = []
        for x in range(0, 470, 10):
            ten_mile_markers_for_graph.append(x+5)
            temp_df = df.loc[df['Specific Sign']==this_sign]
            temp_df = temp_df.loc[temp_df['Mile Post (ten miles)']==x]
            tot_signs = temp_df['Sign Type'].count()
            this_list.append(tot_signs)

        dataset = np.array(this_list)
        if count == 1:
            p1 = plt.bar(ten_mile_markers_for_graph, dataset, bottom = last_dataset, width=10, label = this_sign,
                        hatch = hatch,edgecolor=".2", linewidth=1.5, color = "midnightblue")
        elif count == len(mylist):
            p1 = plt.bar(ten_mile_markers_for_graph, dataset, bottom = last_dataset, width=10, label = this_sign,
                        hatch = hatch,edgecolor=".2", linewidth=1.5, color = "white")
        else:
            p1 = plt.bar(ten_mile_markers_for_graph, dataset, bottom = last_dataset, width=10, label = this_sign,
                        hatch = hatch,edgecolor=".2", linewidth=1.5)
        last_dataset = dataset + last_dataset
        count +=1
        if hatch == "":
            hatch = "/"
        else:
            hatch =""

    plt.legend(fontsize=20, ncol=2)

    bar_max = max(last_dataset)*1.05
    plt.bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    plt.bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    plt.bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax.text(279, bar_max/2,
               '  No Data\n  Collected\n  MP 276-315', fontsize=13,
           bbox=bbox_props)


    ax.text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=15)

    ax.text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=15)
    ax.text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=15)
    ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax.text(175, bar_max*1.15, 'Virginia', fontsize=25)
    ax.text(228, bar_max*1.15, 'North Carolina', fontsize=25)

    #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
    #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax.set_ylim([0, bar_max*1.5])
    ax.set_xlim([0, 470])
#    ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
    ax.set_title(" ",fontsize= 12, pad = 20) # title of plot

    ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax.set_ylabel(axis_title, fontsize = 30)#ylabel
    ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax.tick_params(axis='x', which='major', labelsize=16)
    ax.tick_params(axis='y', which='major', labelsize=16)
    ax.tick_params(axis='both', which='minor', labelsize=8)
    ax.grid('off')
    file_save_name = folder_name + "/" + "Specific Sign_locations.png"
    #print(file_save_name)
    plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
    plt.close()



def make_double_location_plot(sign_df, crash_df, sign, crash, folder_name):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    graph_title = sign + " Sign and " + crash + " Crash" + " Locations on BLRI Parkway"
    axis_title = "Sign and Crash Count"
    plt.style.use('seaborn')
    fig, ax = plt.subplots(1,1,figsize=(20,10))
    tick_spacing = 50

    plt.bar(sign_df['Mile Post (ten miles)'].value_counts().index+5,
            sign_df['Mile Post (ten miles)'].value_counts().values,
            width = 10, label = sign + " Signs")

    plt.bar(crash_df['Mile Post (ten miles)'].value_counts().index+5,
            -crash_df['Mile Post (ten miles)'].value_counts().values,
            width = 10, label = crash + " Crashes")


    bar_max = 1.15*max(sign_df['Mile Post (ten miles)'].value_counts().max(),
                    crash_df['Mile Post (ten miles)'].value_counts().max())

    plt.bar(115, [-bar_max/1.05, bar_max/1.05], width = 30, color = 'gray', alpha = 0.15)
    plt.bar(290, [-bar_max/1.05, bar_max/1.05], width = 20, color = 'gray', alpha = 0.15)
    plt.bar(385, [-bar_max/1.05, bar_max/1.05], width = 30, color = 'gray', alpha = 0.15)
    ax.vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    ax.hlines([0], xmin = 0, xmax = 470,
            linestyles = "solid", linewidth = 3, color = 'black')
            
    plt.bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax.text(279, bar_max/2,
               ' No Sign Data\n Collected\n MP 276-315', fontsize=13,
           bbox=bbox_props)


    ax.text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=15)

    ax.text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=15)
    ax.text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=15)
    ax.vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax.vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax.text(175, bar_max*1.2, 'Virginia', fontsize=25)
    ax.text(228, bar_max*1.2, 'North Carolina', fontsize=25)
    #ax.text(200, bar_max*1.15, 'VA', fontsize=25)
    #ax.text(225, bar_max*1.15, 'NC', fontsize=25)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax.set_ylim([bar_max*-1.3, bar_max*1.6])
    ax.set_xlim([0, 460])
#    ax.set_title(graph_title,fontsize= 35, pad = 20) # title of plot
    ax.set_title(" ",fontsize= 12, pad = 20) # title of plot

    ax.set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax.set_ylabel(axis_title, fontsize = 30)#ylabel
    ax.xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax.tick_params(axis='x', which='major', labelsize=16)
    ax.tick_params(axis='y', which='major', labelsize=16)
    ax.tick_params(axis='both', which='minor', labelsize=8)
    ax.grid('off')
    ax.legend(fontsize = 20, frameon=True,fancybox = True,shadow=True, handlelength=1.5,framealpha=1, )
    file_save_name = folder_name + "/sign_and_crash_locations.png"
    #print(file_save_name)
    plt.savefig(file_save_name, bbox_inches = 'tight', pad_inches = 0)
    plt.close()


def clean_crash_data(df):
    df_clean = df.loc[df['Mile Post (ten miles)']>-10]
    return df_clean

def join_data(crash_df, sign_df):
    import pandas as pd
    import math
    mile_marker = []
    ten_mile_marker = []
    dead_on_crashes = []
    nearby_crashes = []
    total_crashes = []
    lon_list = []
    lat_list = []

    dead_on_signs = []
    nearby_signs = []
    total_signs=[]
    crash_rates = []
    lon, lat = 0, 0
    last_mile = -1

    for x in crash_df['Milepost'].unique():
        if x == last_mile + 0.1:
            continue
        if x==last_mile + 0.2:
            continue
        try:
            ten_mile = math.floor(x/10)*10
        except:
            continue
        mile_marker.append(x)
        ten_mile_marker.append(ten_mile)

        #filter for close crashes
        temp_df = crash_df.loc[crash_df['Milepost']==x]
        dead_on = len(temp_df)
        dead_on_crashes.append(dead_on)

        temp_df = crash_df.loc[(crash_df['Milepost']==x+0.1)|(crash_df['Milepost']==x+0.2)]
        nearby = len(temp_df)
        nearby_crashes.append(nearby)
        total_crash = nearby+dead_on
        total_crashes.append(total_crash)

        #filter for close signs
        temp_df = sign_df.loc[sign_df['Mile Post (tenth of a mile)']==x]
        if len(temp_df)>0:
            lat = temp_df['LAT'].median()
            lon = temp_df['LON'].median()
        dead_on = len(temp_df)
        dead_on_signs.append(dead_on)
        temp_df = sign_df.loc[(sign_df['Mile Post (tenth of a mile)']==x+0.1)|(sign_df['Mile Post (tenth of a mile)']==x-0.1)]
        if (len(temp_df)>0)&(lat==0):
            lat = temp_df['LAT'].median()
            lon = temp_df['LON'].median()
        nearby = len(temp_df)
        nearby_signs.append(nearby)
        total_sign = nearby+dead_on
        total_signs.append(total_sign)

        try:
            crash_rate = total_crash/ total_sign
        except:
            crash_rate = 0

        crash_rates.append(crash_rate)
        lat_list.append(lat)
        lon_list.append(lon)

        last_mile = x
        lat =0
        lon =0

    crash_df_joined = pd.DataFrame({
        "Mile Post (tenth of a mile)": mile_marker,
        "Mile Post (ten miles)": ten_mile_marker,
        "Stop crashes at this location": dead_on_crashes,
        "Stop crashes within tenth of mile": nearby_crashes,
        'Total crashes':total_crashes,
        "Stop signs at this location": dead_on_signs,
        "Stop signs within tenth of mile": nearby_signs,
        'Total signs':total_signs,
        'Crash rate':crash_rates,
        'LON': lon_list,
        'LAT': lat_list
    })
    crash_df_joined = crash_df_joined.loc[crash_df_joined['Total signs']>0]
    return crash_df_joined

def make_milepost_df():
    from bs4 import BeautifulSoup
    x=0

    name_list = []
    coord_list = []
    long_list = []
    lat_list = []
    y=0
    inputfile = "Mileposts.kml"
    with open(inputfile, 'r') as f:
        soup = BeautifulSoup(f)

      # After you have a soup object, you can access tags very easily.
      # For instance, you can iterate over and get <description> like so:

    #    for node in soup.select('Folder'):
    #
    #        print('done')
    #        x+=1
        for node in soup.select('coordinates'):
            name_list.append(node.parent.parent.find('name').contents[0])
            coord = node.contents[0]
            lat = coord[coord.find(",")+1:coord.rfind(",")]
            long = coord[:coord.find(",")]
            long_list.append(long)
            lat_list.append(lat)
            coord_list.append(node.contents[0])
            y+=1

    import pandas as pd
    map_df = pd.DataFrame({'Mile Post':name_list,
                          'LON':long_list,
                          'LAT':lat_list})

    map_df['Mile Post']=map_df['Mile Post'].astype(int)
    map_df['LON']=map_df['LON'].astype(float)
    map_df['LAT']=map_df['LAT'].astype(float)
    map_df = map_df.drop_duplicates()
    map_df = map_df.sort_values(by = ['Mile Post'])

    return map_df


def make_bubble_map(crash_df_joined, sign_name, folder):
    import folium
    import os
    import time
    from selenium import webdriver

    chromedriver_location=r'C:\Users\eric.englin\Downloads/chromedriver.exe'
    # Make an empty map
    m = folium.Map(location=[36.652541846246095, -80.6166736005195],
                   tiles="openstreetmap",
                   zoom_start=7)

    # add bubbles to the map with radius of 5000 x # of crashes
    for i in range(0,len(crash_df_joined)):
        folium.Circle(
            location=[crash_df_joined.iloc[i]['LAT'], crash_df_joined.iloc[i]['LON']],
            popup="Milepost: {}, crashes: {}, signs: {}".format(crash_df_joined.iloc[i]['Mile Post (tenth of a mile)'],
                                                       crash_df_joined.iloc[i]['Total crashes'],
                                                       crash_df_joined.iloc[i]['Total signs']),
            radius=crash_df_joined.iloc[i]['Total crashes']*5000,
            color='crimson',
            fill=True,
            fill_color='crimson'
        ).add_to(m)
        # Save it as html

    locations = crash_df_joined[['LAT', 'LON']]
    locationlist = locations.values.tolist()

    #add point to show exact location of each crash
    for i in range(0, len(locationlist)):
        folium.CircleMarker(location=[crash_df_joined.iloc[i]['LAT'], crash_df_joined.iloc[i]['LON']],
                            radius=1,
                            weight=3,
                           color = 'black').add_to(m)


    #add line of BLRI Parkway
    map_df = make_milepost_df()
    milepost_locations = map_df[['LAT', 'LON']]
    milepost_locationlist = milepost_locations.values.tolist()
    folium.PolyLine(milepost_locationlist, color = "black").add_to(m)


    x=sign_name+ " Signs"
    title =  "./" + folder +  "/" + sign_name + "_sign.html"
    title_html = '''
             <h1 align="center" style="font-size:16px"><b>Sign Type: {}</b></h1>
             '''.format(str(x))

    m.get_root().html.add_child(folium.Element(title_html))
    #print(os.getcwd())
    m.save(title)
    tmpurl='file://{path}/{mapfile}'.format(path=os.getcwd(),mapfile=title)
    #print(tmpurl)
    browser = webdriver.Chrome(chromedriver_location)
    browser.get(tmpurl)
    #Give the map tiles some time to load
    time.sleep(5)
    png_title = folder + "/bubble_map.png"
    browser.save_screenshot(png_title)

    browser.quit()


def simplify_crash_df(df):
    import pandas as pd
    ten_mile_markers = []
    ten_mile_markers_for_graph = []
    ten_mile_crashes = []
    ten_mile_signs = []

    for x in range(0, 470, 10):
        ten_mile_markers.append(x)
        temp_df = df.loc[df['Mile Post (ten miles)']==x]
        dead_on = temp_df['Total crashes'].sum()
        ten_mile_crashes.append(dead_on)
        ten_mile_markers_for_graph.append(x+5)
        #filter for close signs
        dead_on = temp_df['Total signs'].sum()
        ten_mile_signs.append(dead_on)

    df_simple = pd.DataFrame({
        "Mile Post (ten miles)": ten_mile_markers,
        "Graphing mile post":ten_mile_markers_for_graph,
        "Stop crashes": ten_mile_crashes,
        "Stop signs": ten_mile_signs,
    })
    return df_simple, ten_mile_markers, ten_mile_crashes, ten_mile_signs, ten_mile_markers_for_graph

def make_sign_crash_plots(df_simple, traffic_df, sign_name, folder):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    df_simple, ten_mile_markers, ten_mile_crashes, ten_mile_signs, ten_mile_markers_for_graph = simplify_crash_df(df_simple)
    df_simple = df_simple.merge(traffic_df)
    df_simple['Traffic Adjusted Crash Rate'] = df_simple['Stop crashes']/ df_simple['AADT']*1000

    graph_title_total= 'Total Number of Incidents Within 0.2 Miles of ' + sign_name + ' Signs'
    graph_title_traffic= 'Traffic Adjusted Incidents Within 0.2 Miles of ' + sign_name + ' Signs'
    plt.style.use('seaborn')
    fig, ax = plt.subplots(1,2,figsize=(45,15))
    tick_spacing = 50

    #plot data
    bar_max = max(ten_mile_crashes)*1.05

    ax[0].bar(ten_mile_markers_for_graph, ten_mile_crashes,width = 10, label ="Crashes")
    ax[0].bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax[0].bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    ax[0].bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax[0].vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    ax[0].bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax[0].text(279, bar_max/2,
               '  No Data\n  Collected\n  MP 276-315', fontsize=18,
           bbox=bbox_props)


    ax[0].text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=25)

    ax[0].text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=25)
    ax[0].text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=25)
    ax[0].vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax[0].vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax[0].text(160, bar_max*1.15, 'Virginia', fontsize=35)
    ax[0].text(228, bar_max*1.15, 'North Carolina', fontsize=35)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax[0].set_title(graph_title_total,fontsize= 35, pad = 20) # title of plot
    ax[0].set_ylim([0, bar_max*1.3])
    ax[0].set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax[0].set_ylabel('Number of incidents', fontsize = 30)#ylabel
    ax[0].xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax[0].tick_params(axis='x', which='major', labelsize=25,rotation = 0)
    ax[0].tick_params(axis='y', which='major', labelsize=25)
    ax[0].tick_params(axis='both', which='minor', labelsize=8)
    ax[0].grid('off')
    #plt.legend(fontsize = 20)

    bar_max = max(df_simple['Traffic Adjusted Crash Rate'])*1.05
    #plot data
    ax[1].bar(ten_mile_markers_for_graph, df_simple['Traffic Adjusted Crash Rate'],width = 10, label ="Crashes")
    ax[1].bar(115, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax[1].bar(290, bar_max/1.05, width = 20, color = 'gray', alpha = 0.15)
    ax[1].bar(385, bar_max/1.05, width = 30, color = 'gray', alpha = 0.15)
    ax[1].vlines([280, 300], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 0.25)

    ax[1].bar(295.5, [-bar_max/1.05, bar_max/1.05], width = 39, color = 'salmon', alpha = 0.25)
    bbox_props = dict(boxstyle="round", fc="w", ec="0.5", alpha=0.9)
    ax[1].text(279, bar_max/2,
               '  No Data\n  Collected\n  MP 276-315', fontsize=18,
           bbox=bbox_props)


    ax[1].text(95, bar_max,
               'Roanoke, VA\n MP 100-130', fontsize=25)

    ax[1].text(275,bar_max,
               'Boone, NC\n MP 280-300', fontsize=25)
    ax[1].text(365, bar_max,
               'Asheville,NC\n MP 370-400', fontsize=25)
    ax[1].vlines([100, 130, 370,400], ymin=-bar_max/1.05, ymax = bar_max/1.05,
                linestyles = "dotted", linewidth = 2)
    ax[1].vlines([220], ymin=-bar_max*1.2, ymax = bar_max*1.2,
                linestyles = "dotted", linewidth = 5, color = 'black')
    ax[1].text(160, bar_max*1.15, 'Virginia', fontsize=35)
    ax[1].text(228, bar_max*1.15, 'North Carolina', fontsize=35)
    #ax.bar(ten_mile_markers_for_graph, ten_mile_signs,width = 5,label ="Stop Signs")

    #make graph pretty
    ax[1].set_title(graph_title_traffic,fontsize= 35, pad = 20) # title of plot
    ax[1].set_ylim([0,bar_max*1.3])
    ax[1].set_xlabel('Milepost (10-mile segments)',fontsize = 30) #xlabel
    ax[1].set_ylabel('Number of incidents per 1,000 vehicles', fontsize = 30)#ylabel
    ax[1].xaxis.set_major_locator(ticker.MultipleLocator(tick_spacing))
    ax[1].tick_params(axis='x', which='major', labelsize=25,rotation = 0)
    ax[1].tick_params(axis='y', which='major', labelsize=25)
    ax[1].tick_params(axis='both', which='minor', labelsize=8)
    ax[1].grid('off')
    #plt.legend(fontsize = 20)
    save_file_location = "./" + folder + "/crashes_near_signs.png"
    plt.savefig(save_file_location, bbox_inches = 'tight', pad_inches = 0)
    plt.close()
    #return plt

def make_sign_bar_chart(df, field, sign_name):
    #same function as make_sign_charts but no pie chart
    import matplotlib.pyplot as plt
    import pandas as pd

    standard_response_length = {"Multiple Signs":2,
                           "Parkway Location":2,
                           "Post Material":3,
                           "Sign Condition":3,
                           "Sign Owner":2,
                           "Sign Face":6,
                           "Number of posts":2,
                           "Sign Material":6,
                           "Physical Location":6,
                           "Maintenance Required":6,
                           "Post Size":6,
                           'District':4,
                           "Sign Type":7,
                           "Sign Subtype":10,
                           "Specific Sign":10}

    num_standard_responses = standard_response_length[field]

    plt.style.use('seaborn')
    x=df[field].value_counts().index
    y = df[field].value_counts()
    #print(field, ": ", len(y))
    percent = 100.*y/y.sum()
    max_y =  max(y)
    other_options = len(y)-num_standard_responses
    other_count = y[num_standard_responses:].sum()
    other_percent = percent[num_standard_responses:].sum()
    x=list(df[field].value_counts().index.astype(str))
    y = list(y)
    percent = list(percent)
    if len(y)>num_standard_responses:
        if field == "Sign Subtype" or field == "Specific Sign":
            y = y[:num_standard_responses]
            percent = percent[:num_standard_responses]
            x = x[:num_standard_responses]
        else:
            y = y[:num_standard_responses]
            y.append(other_count)
            percent = percent[:num_standard_responses]
            percent.append(other_percent)
            x = x[:num_standard_responses]
            x.append("Other Response")


    text_list = []
    for tot_count, perc in zip(y, percent):
        new_count = '{:,.0f} '.format(tot_count)
        perc = '({:.0f}%)'.format(perc)
        this_sign_text = new_count + perc
        text_list.append(this_sign_text)

    fig, ax = plt.subplots(1,1,figsize=(30,15))

    ax.barh(x, width = y)
    ax.set_title(field + ", " + str(other_count) + " signs excluded, " + str(other_options) + "categories excluded",fontsize= 35)
    ax.set_xlabel('Number of signs',fontsize = 30) #xlabel
    ax.set_ylabel('', fontsize = 30)#ylabel
    ax.set_xlim([0,max_y*1.15])
    ax.tick_params(axis='x', which='major', labelsize=25,rotation = 0)
    ax.tick_params(axis='y', which='major', labelsize=25)
    ax.grid('off')
    for p, text in zip(ax.patches, text_list):
        percentage ='{:,.0f}KM²'.format(p.get_width())
        width, height =p.get_width(),p.get_height()
        this_x=p.get_x()+width+(max_y*0.025)
        this_y=p.get_y()+height/2.5
        ax.annotate(text,(this_x,this_y),size=25)

    save_file_location = "./" + sign_name + "/" + field + "_combo.png"
    plt.savefig(save_file_location, bbox_inches = 'tight', pad_inches = 0)
    plt.close()

def make_sign_charts(df, field, sign_name):
    #modified from: https://stackoverflow.com/questions/23577505/how-to-avoid-overlapping-of-labels-autopct-in-a-matplotlib-pie-chart
    import matplotlib.pyplot as plt
    import pandas as pd

    standard_response_length = {"Multiple Signs":2,
                           "Parkway Location":2,
                           "Post Material":3,
                           "Sign Condition":3,
                           "Sign Owner":2,
                           "Sign Face":6,
                           "Number of posts":2,
                           "Sign Material":6,
                           "Physical Location":6,
                           "Maintenance Required":6,
                           "Post Size":6,
                           'District':4}

    num_standard_responses = standard_response_length[field]

    plt.style.use('seaborn')
    x=df[field].value_counts().index
    y = df[field].value_counts()
    #print(field, ": ", len(y))
    percent = 100.*y/y.sum()
    other_count = y[num_standard_responses:].sum()
    other_percent = percent[num_standard_responses:].sum()
    x=list(df[field].value_counts().index.astype(str))
    y = list(y)
    percent = list(percent)
    if len(y)>num_standard_responses:
        y = y[:num_standard_responses]
        y.append(other_count)
        percent = percent[:num_standard_responses]
        percent.append(other_percent)
        x = x[:num_standard_responses]
        x.append("Other Response")

    explode_list = [0]* len(x)
    explode_list[0] = 0.1
    fig, ax = plt.subplots(1,2,figsize=(45,15))
    patches, texts = ax[0].pie(y,
            startangle=90, explode = explode_list)

    labels = ['{0} - {1:1.2f} %'.format(i,j) for i,j in zip(x, percent)]
    ax[0].legend( labels = labels,loc='upper right', bbox_to_anchor=(-0.1, 1.),
               fontsize=30)
    ax[0].set_title(field,fontsize= 40)
#    save_file_location = "./" + sign_name + "/" + field + "_pie.png"
    #ax1.figure.savefig(save_file_location, bbox_inches = 'tight', pad_inches = 0)

#    fig2, ax2 = plt.subplots()
    ax[1].barh(x, width = y)
    ax[1].set_title(field,fontsize= 35)
    ax[1].set_xlabel('Number of signs',fontsize = 30) #xlabel
    ax[1].set_ylabel('', fontsize = 30)#ylabel
    ax[1].tick_params(axis='x', which='major', labelsize=25,rotation = 0)
    ax[1].tick_params(axis='y', which='major', labelsize=25)
    ax[1].grid('off')
#    save_file_location = "./" + sign_name + "/" + field + "_bar.png"
    save_file_location = "./" + sign_name + "/" + field + "_combo.png"
    plt.savefig(save_file_location, bbox_inches = 'tight', pad_inches = 0)
    plt.close()
    #return plt

def write_document(folder, sign_name, number_of_signs, number_of_crashes, location_with_crash,signs_next_to_crash, crashes_next_to_signs,no_maintenance_required_percent, multiple_signs_percent, nps_owner_percent, good_sign_percent, Scotchlite_percent):
    #write word document
    import docx
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
    import os
    try:
        doc = docx.Document("./" + folder + "/_overview.docx")
    except:
        doc = docx.Document("blank.docx") # connect test.docx file
    new_directory = "./" + folder + "/"
    os.chdir(new_directory)
    styles = doc.styles
    style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph(sign_name + " Analysis", style = "Heading 1")
    doc.add_paragraph("Filters:", style = "Heading 2")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("Remove duplicates using: mile post tenth, side of road, road-type, and specific sign")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("Remove all signs more than 1 mile away from road")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("Link crashes to signs within 0.2 miles")
    doc.add_paragraph("Key data points:", style = "Heading 2")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("# of signs (no duplicates): " + str(number_of_signs))
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("# of crashes considered: " + str(number_of_crashes))
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("# of locations with a crash next to a sign: " + str(location_with_crash))
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("# of signs next to a crash: " + str(signs_next_to_crash))
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("# of crashes next to signs: " + str(crashes_next_to_signs))

    #start of sign section
    doc.add_paragraph("\n\n Key Sign Figures:", style = "Heading 2")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("% of signs that don't require maintenance: " + str(no_maintenance_required_percent) + "%")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("% of signs part of multiple signs: " + str(multiple_signs_percent) + "%")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("% of signs that owned by NPS: " + str(nps_owner_percent) + "%")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("% of signs in good condition: " + str(good_sign_percent) + "%")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    r = p.add_run()
    r.add_text("% of signs with Scotchlite sign face: " + str(Scotchlite_percent) + "%")
    p = doc.add_paragraph()
    p.style = 'List Paragraph'
    doc.add_picture("Sign_locations.png", width=Inches(6.9))
    doc.add_picture("Sign Type_locations.png", width=Inches(6.9))
    doc.add_picture("Sign Type All Sequential_locations.png", width=Inches(6.9))
    doc.add_picture("Sign Type_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Subtype_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Warning Subtypes_locations.png", width=Inches(6.9))
    doc.add_picture("Sign Regulatory Subtypes_locations.png", width=Inches(6.9))
    doc.add_picture("Specific Sign_combo.png", width=Inches(6.9))
    doc.add_picture("Specific Sign_locations.png", width=Inches(6.9))


    doc.add_paragraph("\n\n Key Crash Figures:", style = "Heading 2")
    #doc.add_picture("bubble_map.png", width=Inches(6.9))
    doc.add_picture("Crash_locations.png", width=Inches(6.9))
    doc.add_picture("crashes_near_signs.png", width=Inches(6.9))
    doc.add_picture("sign_and_crash_locations.png", width=Inches(6.9))
    #Sign subgraphs
    doc.add_paragraph("\n\n Other Sign Breakout Graphs:", style = "Heading 2")
    doc.add_picture("Number of posts_combo.png", width=Inches(6.9))
    doc.add_picture("Post Material_combo.png", width=Inches(6.9))
    doc.add_picture("Post Size_combo.png", width=Inches(6.9))
    doc.add_picture("Maintenance Required_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Material_combo.png", width=Inches(6.9))
    doc.add_picture("District_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Condition_combo.png", width=Inches(6.9))
    doc.add_picture("Parkway Location_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Owner_combo.png", width=Inches(6.9))
    doc.add_picture("Multiple Signs_combo.png", width=Inches(6.9))
    doc.add_picture("Sign Face_combo.png", width=Inches(6.9))
    doc.add_picture("Physical Location_combo.png", width=Inches(6.9))

    file_save_name = folder + "_overview.docx"
    doc.save(file_save_name)
    os.chdir("C:\\Users\\eric.englin\\Desktop\\BLRI\\analysis\\")


def do_whole_sign_analysis(sign_df, crash_df, traffic_df, sign, crash, folder):

    number_of_signs = len(sign_df)
    number_of_crashes = len(crash_df)
    no_maintenance_required_percent = round(len(sign_df.loc[sign_df['Maintenance Required']=="None"])/len(sign_df)*100,1)
    multiple_signs_percent = round(len(sign_df.loc[sign_df['Multiple Signs']=="Yes"])/len(sign_df)*100,1)
    nps_owner_percent = round(len(sign_df.loc[sign_df['Sign Owner']=="NPS"])/len(sign_df)*100,1)
    good_sign_percent = round(len(sign_df.loc[sign_df['Sign Condition']=="Good"])/len(sign_df)*100,1)
    Scotchlite_percent = round(len(sign_df.loc[sign_df['Sign Face'].isin(['Scotchlite Reflective',
                                                               'Scotchlite'])])/len(sign_df)*100,1)

    for x in ['Sign Face','Multiple Signs', 'Sign Owner','Parkway Location', 'Physical Location',
           'District', 'Sign Condition', 'Maintenance Required', 'Sign Material',
              'Post Material', 'Post Size', 'Number of posts',"Sign Type",
              'Sign Subtype','Specific Sign']:
        plt = make_sign_bar_chart(sign_df, x, folder)


    plt = make_location_plot(sign_df, "Sign", sign, folder)
    plt = make_location_plot(crash_df, "Crash", crash,folder)
    plt = make_double_location_plot(sign_df, crash_df, sign, crash, folder)
    plt = make_location_subtype_plots(sign_df, "Sign", sign, folder)
    plt = make_location_specific_sign_plots(sign_df, "Sign", sign, folder)
    crash_df_joined = join_data(crash_df, sign_df)
    location_with_crash = len(crash_df_joined)
    signs_next_to_crash = crash_df_joined['Total signs'].sum()
    crashes_next_to_signs = crash_df_joined['Total crashes'].sum()
    crash_df_joined2 = crash_df_joined.reset_index()
    #make_bubble_map(crash_df_joined2,  sign, folder)
    plt = make_sign_crash_plots(crash_df_joined2, traffic_df, sign, folder)
    write_document(folder, sign, number_of_signs, number_of_crashes, location_with_crash,
              signs_next_to_crash, crashes_next_to_signs,
              no_maintenance_required_percent, multiple_signs_percent,
               nps_owner_percent, good_sign_percent, Scotchlite_percent)
